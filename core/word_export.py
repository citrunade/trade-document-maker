"""
Word (.docx) 单据导出模块
使用 python-docx 生成单据，字段内容与 PDF 导出（core/pdf_export.py）保持一致。
Own（本公司）与收件方信息放入页眉、银行基本信息放入页脚，随 Word 原生页眉页脚机制
自动重复显示在每一页，效果与 PDF 版式一致。
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from core import calc

DOC_TITLES = {
    "PI": ("PROFORMA INVOICE", "形式发票"),
    "CI": ("COMMERCIAL INVOICE", "商业发票"),
    "PL": ("PACKING LIST", "装箱单"),
}

CJK_FONT = "Microsoft YaHei"
HEADER_BLUE = RGBColor(0x5B, 0x9B, 0xD5)


def _set_cjk(run, font_name=CJK_FONT, size=None, bold=None, color=None):
    """同时设置西文与中文字体，避免 Word 默认字体不支持中文导致乱码/方块"""
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def _para(container, text="", size=9, bold=False, align=None, color=None):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_cjk(run, size=size, bold=bold, color=color)
    return p


def _cell_text(cell, text="", size=9, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    _set_cjk(run, size=size, bold=bold, color=color)
    return p


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color})
    tcPr.append(shd)


def export_document(doc: dict, filepath: str) -> str:
    doc_type = doc.get("doc_type", "PI")
    financial = doc_type != "PL"
    currency = doc.get("currency", "USD")
    customer = doc.get("customer_snapshot", {})
    own = doc.get("own_snapshot", {})
    delivery = doc.get("delivery_snapshot", {})
    conditions = doc.get("conditions_snapshot", {})
    banking = doc.get("banking_snapshot", {})
    title_en, title_cn = DOC_TITLES.get(doc_type, DOC_TITLES["PI"])

    document = Document()
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(1.5)
    section.top_margin = section.bottom_margin = Cm(1.5)

    # ---------------- 页眉：左收件方 / 右 Own（重复显示在每一页） ----------------
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=section.page_width - section.left_margin - section.right_margin)
    header_table.autofit = True
    left_cell, right_cell = header_table.rows[0].cells

    left_cell.text = ""
    _p = left_cell.paragraphs[0]
    _set_cjk(_p.add_run(customer.get("name_en", "") + "\n"), size=10, bold=True)
    for text in (customer.get("name_cn", ""), customer.get("address_en", "")):
        if text:
            _p.add_run("\n")
            _set_cjk(_p.add_run(text), size=8.5)
    for label, value in (("Phone", customer.get("tel_phone", "")),
                          ("Company Reg. No.", customer.get("company_reg_no", "")),
                          ("GST No.", customer.get("gst_no", ""))):
        if value:
            _p.add_run("\n")
            _set_cjk(_p.add_run(f"{label}: {value}"), size=8.5)

    right_cell.text = ""
    _p2 = right_cell.paragraphs[0]
    _p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_cjk(_p2.add_run(f"{title_en}\n"), size=13, bold=True, color=HEADER_BLUE)
    _set_cjk(_p2.add_run(f"{title_cn}\n"), size=9)
    _set_cjk(_p2.add_run(own.get("company_name_en", "") + "\n"), size=10, bold=True)
    for text in (own.get("company_name_cn", ""),):
        if text:
            _p2.add_run("\n")
            _set_cjk(_p2.add_run(text), size=8.5)
    for line in (own.get("address", "") or "").splitlines():
        if line.strip():
            _p2.add_run("\n")
            _set_cjk(_p2.add_run(line.strip()), size=8.5)
    for label, value in (("Phone", own.get("phone", "")),
                          ("Company Reg. No.", own.get("company_reg_no", "")),
                          ("GST No.", own.get("gst_no", ""))):
        if value:
            _p2.add_run("\n")
            _set_cjk(_p2.add_run(f"{label}: {value}"), size=8.5)

    # ---------------- 页脚：银行基本信息（重复显示在每一页） ----------------
    footer = section.footer
    _p3 = footer.paragraphs[0]
    _p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cjk(_p3.add_run(own.get("company_name_en", "") + "\n"), size=8, bold=True)
    bank_parts = []
    if banking.get("bank_name"):
        bank_parts.append(f"Bank: {banking.get('bank_name')}")
    if banking.get("account_no"):
        bank_parts.append(f"Account No.: {banking.get('account_no')}")
    if banking.get("swift_code"):
        bank_parts.append(f"Swift Code: {banking.get('swift_code')}")
    _set_cjk(_p3.add_run("   ".join(bank_parts)), size=7.5, color=RGBColor(0x80, 0x80, 0x80))

    # ---------------- 正文 ----------------
    _para(document, f"{title_en} / {title_cn}", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    document.add_paragraph()

    incoterm_line = conditions.get("incoterm", "")
    if incoterm_line and doc.get("destination"):
        incoterm_line = f"{incoterm_line}  {doc.get('destination')}"

    info_table = document.add_table(rows=1, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left_col, right_col = info_table.rows[0].cells

    def _box(cell, title, rows):
        cell.text = ""
        _cell_text(cell, title, size=9.5, bold=True)
        _shade_cell(cell, "DCE9F7")
        for label, value in rows:
            p = cell.add_paragraph()
            if not label:
                continue
            run_label = p.add_run(f"{label}  ")
            _set_cjk(run_label, size=8.5, bold=True)
            _set_cjk(p.add_run(str(value)), size=8.5)

    _box(left_col, "Delivery Address / Conditions", [
        (delivery.get("company_name", ""), ""),
        ("", delivery.get("address", "")),
        ("Terms of Payment", conditions.get("terms_of_payment", "")),
        ("Incoterms", incoterm_line),
        ("Shipment by", conditions.get("shipment_by", "")),
    ])
    _box(right_col, "Information", [
        ("Doc No", doc.get("doc_number", "")),
        ("Document Date", doc.get("date", "")),
        ("Validity Start Date", doc.get("validity_start", "")),
        ("Validity End Date", doc.get("validity_end", "")),
        ("Our Contact", own.get("contact_person", "")),
        ("Telephone", own.get("telephone", "")),
        ("Email", own.get("email", "")),
    ])
    document.add_paragraph()

    # ---------------- 明细表 ----------------
    totals = calc.compute_totals(doc.get("lines", []))
    if financial:
        headers = ["No.", "Description", "Qty", "Unit", "Unit Price", "Total Price",
                   "COO", "Net Weight", "Total N.W.", "HS Code", "Remark"]
    else:
        headers = ["No.", "Description", "Qty", "Unit", "COO", "Net Weight", "Total N.W.", "HS Code", "Remark"]

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _cell_text(cell, h, size=8.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(cell, "5B9BD5")

    for i, line in enumerate(totals["lines"], start=1):
        row_cells = table.add_row().cells
        desc = line.get("name_en", "")
        if line.get("name_cn"):
            desc += f"\n{line.get('name_cn')}"
        values = [str(i), desc, f"{line.get('quantity', 0):g}", line.get("unit", "")]
        if financial:
            values += [f"{line.get('unit_price', 0):,.2f}", f"{line.get('subtotal', 0):,.2f}"]
        values += [line.get("coo", ""), f"{line.get('net_weight', 0):.2f}",
                   f"{line['total_net_weight']:.2f}", line.get("hs_code", ""), line.get("remark", "")]
        for c, val in enumerate(values):
            _cell_text(row_cells[c], val, size=8)

    document.add_paragraph()
    if financial:
        _para(document, f"Total Amount: {currency} {totals['total_amount']:,.2f}", size=9.5, bold=True,
              align=WD_ALIGN_PARAGRAPH.RIGHT)
        _para(document, calc.amount_in_words(totals["total_amount"], currency), size=8.5, bold=True)

    _para(document, f"Total Qty: {totals['total_quantity']:g}    Total N.W.: {totals['total_net_weight']:.2f} kg", size=9)
    document.add_paragraph()

    if doc.get("remark"):
        _para(document, "Note / 备注", size=9.5, bold=True)
        _para(document, doc.get("remark", ""), size=8.5)
        document.add_paragraph()

    beneficiary_rows = [
        ("Beneficiary Name", banking.get("beneficiary_name", "")),
        ("Beneficiary Bank Name", banking.get("beneficiary_bank_name", "")),
        ("Swift Code", banking.get("beneficiary_swift_code", "")),
        ("Beneficiary Bank's Correspondent Bank", banking.get("correspondent_bank", "")),
        ("Beneficiary Bank Address", banking.get("beneficiary_bank_address", "")),
        ("Beneficiary Account No.", banking.get("beneficiary_account_no", "")),
    ]
    if any(v for _, v in beneficiary_rows):
        _para(document, "Beneficiary Bank Details / 收款银行详情", size=9.5, bold=True)
        for label, value in beneficiary_rows:
            if not value:
                continue
            p = document.add_paragraph()
            _set_cjk(p.add_run(f"{label}: "), size=8.5, bold=True)
            _set_cjk(p.add_run(str(value)), size=8.5)

    document.save(filepath)
    return filepath
