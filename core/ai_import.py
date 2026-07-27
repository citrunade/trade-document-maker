"""
AI 文件导入模块
从 PDF / Excel / Word / 图片 (JPG/PNG) 中提取客户信息、产品信息或单据明细，
通过阿里云百炼（Qwen）完成 OCR 与结构化字段抽取。

处理策略：
- Excel：直接读取表格文字内容（本身已是结构化数据，无需 OCR）。若表格文字内容过于稀疏
  （例如产品信息以图片形式贴在单元格中，而非文字），自动改用表格中嵌入的第一张图片调用视觉模型。
- Word (.docx)：提取段落与表格文字；若文字内容过于稀疏（例如整页以图片形式排版），
  自动改用文档中嵌入的第一张图片调用视觉模型。
- PDF：优先提取页面中的文字层（免费、快速）；如果某页没有文字层（如扫描件/图片型 PDF），
  才将该页渲染为图片调用视觉模型（费用更高，仅在必要时使用）。
- JPG/PNG：直接调用视觉模型。

所有函数返回结构化 dict/list，调用方（UI 层）负责展示审核对话框，
本模块不直接写入任何本地数据文件。
"""
import json
import os
import re

import fitz  # PyMuPDF

from core import ai_client
from core.models import make_customer, make_product, make_doc_line

SUPPORTED_EXTENSIONS = {".pdf", ".xlsx", ".xls", ".docx", ".jpg", ".jpeg", ".png"}

# 判定"文字内容过于稀疏、需改用图片识别"的最小有效字符数阈值
MIN_TEXT_CHARS = 20


class AIImportError(Exception):
    pass


# ---------------- 文件类型识别 ----------------
def get_file_kind(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return "pdf"
    if ext in (".xlsx", ".xls"):
        return "xlsx"
    if ext == ".docx":
        return "docx"
    if ext == ".doc":
        raise AIImportError(
            "不支持旧版 .doc 格式，请在 Word 中另存为 .docx 格式后重新导入。"
        )
    if ext in (".jpg", ".jpeg", ".png"):
        return "image"
    raise AIImportError(f"不支持的文件类型：{ext}（仅支持 PDF / Excel / Word(.docx) / JPG / PNG）")


# ---------------- PDF 内容提取 ----------------
def extract_pdf_pages(path: str) -> list:
    """
    逐页提取 PDF 内容。
    返回列表，每项为 ("text", str) 或 ("image", bytes)：
    - 若该页存在文字层，返回提取到的文本；
    - 若该页没有文字层（判定为扫描件/图片型页面），返回该页渲染后的 PNG 图片字节。
    """
    results = []
    doc = fitz.open(path)
    try:
        for page in doc:
            text = page.get_text().strip()
            if text:
                results.append(("text", text))
            else:
                # 无文字层，判定为扫描件，渲染为图片交给视觉模型处理
                pix = page.get_pixmap(dpi=200)
                results.append(("image", pix.tobytes("png")))
    finally:
        doc.close()
    return results


# ---------------- Excel 内容提取 ----------------
def xlsx_to_markdown(path: str) -> str:
    import pandas as pd
    df = pd.read_excel(path, dtype=str).fillna("")
    return df.to_markdown(index=False)


def xlsx_extract_images(path: str) -> list:
    """
    提取 Excel 文件中直接贴图/嵌入的图片（不含单元格文字），用于表格文字内容
    过于稀疏时的视觉模型回退方案（例如产品目录以图片贴在表格里，而非文字）。
    仅支持 .xlsx（openpyxl），.xls 旧格式不支持嵌入图片提取。
    """
    if os.path.splitext(path)[1].lower() != ".xlsx":
        return []
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path)
    except Exception:
        return []
    images = []
    for ws in wb.worksheets:
        for img in getattr(ws, "_images", []):
            try:
                images.append(img._data())
            except Exception:
                continue
    return images


# ---------------- Word 内容提取 ----------------
def extract_docx_content(path: str) -> tuple:
    """
    提取 .docx 文件的段落与表格文字，以及所有嵌入图片。
    返回 (text, images)：text 为拼接后的纯文字内容，images 为图片字节列表。
    """
    import docx
    document = docx.Document(path)

    parts = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts)

    images = []
    for rel in document.part.rels.values():
        if "image" in rel.reltype:
            try:
                images.append(rel.target_part.blob)
            except Exception:
                continue

    return text, images


# ---------------- JSON 解析辅助 ----------------
def _parse_json_response(raw: str):
    """
    从模型回复中提取 JSON。模型有时会用 ```json ... ``` 包裹，需先去除代码块标记。
    解析失败时抛出 AIImportError，绝不返回猜测性的部分数据。
    """
    text = raw.strip()
    fence_match = re.search(r"```(?:json)?\s*([\s\S]*?)```", text)
    if fence_match:
        text = fence_match.group(1).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError as e:
        raise AIImportError(f"AI 返回的内容不是有效的 JSON，无法自动提取字段：{e}\n原始回复：{raw[:500]}")


# ---------------- 提示词 ----------------
CUSTOMER_PROMPT = """你是外贸单据信息提取助手。用户会提供一份文件内容（发票、箱单、产品目录、名片、公司信笺
或类似文件）。只要识别出具备企业主体特征的信息（如买方、卖方、目录内企业信息等），一律判定为客户主体，
不做买方/卖方角色区分；无法明确区分角色时也不要过滤，尽量捕获入库。
请从中提取客户信息，并仅以如下 JSON 格式返回（找不到的字段填空字符串，不要编造信息；地址保持整段
原文文本，不要拆分街道/门牌号/邮编）：
{
  "name_cn": "中文名称",
  "name_en": "英文名称",
  "country_region": "国家/地区",
  "city": "城市",
  "address_en": "英文完整地址（整段文本）",
  "tax_no": "税号/VAT",
  "contact_person": "联系人",
  "email": "电子邮箱",
  "tel_phone": "联系电话",
  "company_reg_no": "公司注册号",
  "gst_no": "GST/增值税登记号",
  "consignee": "收货人信息（如与地址不同，可多行，用\\n分隔）",
  "notify_party": "通知人信息",
  "pod": "目的港（若单据中出现具体港口名）",
  "remark": "其他备注信息"
}
只返回 JSON，不要包含任何其他说明文字。"""

PRODUCT_PROMPT = """你是外贸单据信息提取助手。用户会提供一份产品规格表/报价单/产品目录文件内容，
请从中提取所有产品条目，并仅以如下 JSON 对象格式返回（每个产品一个对象，找不到的字段填 0 或空字符串，不要编造信息）：
{
  "items": [
    {
    "model_no": "型号",
    "name_cn": "中文品名",
    "name_en": "英文品名",
    "hs_code": "HS编码",
    "unit": "单位（如 pcs/set/m）",
    "net_weight": 0.0,
    "gross_weight": 0.0,
    "length_mm": 0.0,
    "width_mm": 0.0,
    "height_mm": 0.0,
    "unit_price": 0.0,
    "currency": "USD",
    "coo": "原产国/地区代码（如 DE/CN，找不到留空）",
    "remark": ""
    }
  ]
}
只返回有效 JSON 对象，不要包含任何其他说明文字。"""

DOCUMENT_LINES_PROMPT = """你是外贸单据信息提取助手。用户会提供一份采购订单(PO)或类似文件内容，
请从中提取所有产品明细行，并仅以如下 JSON 对象格式返回（每行一个对象，找不到的字段填 0 或空字符串，不要编造信息）：
{
  "items": [
    {
    "model_no": "型号",
    "name_cn": "中文品名",
    "name_en": "英文品名",
    "hs_code": "HS编码",
    "unit": "单位",
    "quantity": 0.0,
    "unit_price": 0.0,
    "amount": 0.0,
    "coo": "原产国/地区代码（如 DE/CN，找不到留空）",
    "remark": "备注（找到才填写）"
    }
  ]
}
字段说明（务必区分，不要混淆）：
- quantity：数量/件数（Qty），即购买的件数，通常是较小的整数或简单数字。
- unit_price：单价（Unit Price），即每一件的价格，不是总价/金额（Amount/Total）。
  如果表格中同时存在「单价」和「金额/总价」两列，金额 = 数量 × 单价，
  请只取「单价」列的值填入 unit_price，绝不能把金额或数量误填入 unit_price，
  也不能把单价误填入 quantity。
- amount：金额/总价（Amount/Total），如表格中有该列，原样提取其数字，不要自己计算；
  找不到该列则填 0。此字段仅用于核对 quantity × unit_price 是否正确，不会直接使用。
只返回有效 JSON 对象，不要包含任何其他说明文字。"""


def _run_extraction(path: str, prompt: str, config: dict) -> str:
    """
    根据文件类型选择合适的处理路径，调用阿里云百炼 API，返回模型的原始文本回复。
    """
    api_key = config.get("bailian_api_key", "")
    base_url = config.get("bailian_base_url") or ai_client.DEFAULT_BASE_URL
    text_model = config.get("bailian_text_model") or ai_client.DEFAULT_TEXT_MODEL
    vision_model = config.get("bailian_vision_model") or ai_client.DEFAULT_VISION_MODEL

    kind = get_file_kind(path)

    if kind == "xlsx":
        table_text = xlsx_to_markdown(path)
        if len(table_text.strip()) >= MIN_TEXT_CHARS:
            return ai_client.extract_from_text(api_key, base_url, text_model, prompt, table_text)
        # 表格文字内容过于稀疏（如产品信息以图片贴在单元格中），改用嵌入图片走视觉模型
        images = xlsx_extract_images(path)
        if images:
            return ai_client.extract_from_image(api_key, base_url, vision_model, prompt, images[0], "image/png")
        if table_text.strip():
            return ai_client.extract_from_text(api_key, base_url, text_model, prompt, table_text)
        raise AIImportError("无法从该 Excel 文件中读取任何可识别内容（表格为空且未找到嵌入图片）。")

    if kind == "docx":
        text, images = extract_docx_content(path)
        if len(text.strip()) >= MIN_TEXT_CHARS:
            return ai_client.extract_from_text(api_key, base_url, text_model, prompt, text)
        # 文字内容过于稀疏（如整页以图片排版），改用嵌入图片走视觉模型
        if images:
            return ai_client.extract_from_image(api_key, base_url, vision_model, prompt, images[0], "image/png")
        if text.strip():
            return ai_client.extract_from_text(api_key, base_url, text_model, prompt, text)
        raise AIImportError("无法从该 Word 文件中读取任何可识别内容（文档为空且未找到嵌入图片）。")

    if kind == "image":
        with open(path, "rb") as f:
            image_bytes = f.read()
        ext = os.path.splitext(path)[1].lower()
        mime = "image/jpeg" if ext in (".jpg", ".jpeg") else "image/png"
        return ai_client.extract_from_image(api_key, base_url, vision_model, prompt, image_bytes, mime)

    if kind == "pdf":
        pages = extract_pdf_pages(path)
        if not pages:
            raise AIImportError("无法从该 PDF 中读取任何内容（可能是空白文件）。")
        # 只处理前 5 页，避免超长文档导致费用过高或超出模型上下文
        pages = pages[:5]
        if all(p[0] == "text" for p in pages):
            combined_text = "\n\n".join(p[1] for p in pages)
            return ai_client.extract_from_text(api_key, base_url, text_model, prompt, combined_text)
        # 存在扫描页，使用视觉模型处理第一张图片页（多页扫描件建议用户拆分导入）
        image_page = next((p for p in pages if p[0] == "image"), None)
        if image_page:
            return ai_client.extract_from_image(api_key, base_url, vision_model, prompt, image_page[1], "image/png")
        combined_text = "\n\n".join(p[1] for p in pages if p[0] == "text")
        return ai_client.extract_from_text(api_key, base_url, text_model, prompt, combined_text)

    raise AIImportError(f"未知文件类型：{kind}")


# ---------------- 对外接口 ----------------
def import_customer(path: str, config: dict) -> dict:
    """从文件中提取客户信息，返回一份符合 make_customer 字段结构的 dict（未保存）"""
    raw = _run_extraction(path, CUSTOMER_PROMPT, config)
    data = _parse_json_response(raw)
    if not isinstance(data, dict):
        raise AIImportError("AI 返回的客户信息格式不正确（应为 JSON 对象）。")
    customer = make_customer()
    for key in (
        "name_cn", "name_en", "country_region", "city", "address_en", "tax_no",
        "contact_person", "email", "tel_phone", "company_reg_no", "gst_no",
        "consignee", "notify_party", "pod", "remark",
    ):
        if key in data and data[key] is not None:
            customer[key] = str(data[key])
    return customer


def import_products(path: str, config: dict) -> list:
    """从文件中提取产品列表，返回符合 make_product 字段结构的 dict 列表（未保存）"""
    raw = _run_extraction(path, PRODUCT_PROMPT, config)
    data = _parse_json_response(raw)
    # 百炼 JSON Mode 顶层使用对象；兼容旧模型曾返回的直接数组格式。
    if isinstance(data, dict):
        data = data.get("items", [])
    if not isinstance(data, list):
        raise AIImportError("AI 返回的产品信息格式不正确（应包含 items 数组）。")
    products = []
    for item in data:
        if not isinstance(item, dict):
            continue
        product = make_product(
            model_no=str(item.get("model_no", "")),
            name_cn=str(item.get("name_cn", "")),
            name_en=str(item.get("name_en", "")),
            hs_code=str(item.get("hs_code", "")),
            unit=str(item.get("unit", "pcs")) or "pcs",
            net_weight=_to_float(item.get("net_weight", 0)),
            gross_weight=_to_float(item.get("gross_weight", 0)),
            length_mm=_to_float(item.get("length_mm", 0)),
            width_mm=_to_float(item.get("width_mm", 0)),
            height_mm=_to_float(item.get("height_mm", 0)),
            unit_price=_to_float(item.get("unit_price", 0)),
            currency=str(item.get("currency", "USD")) or "USD",
            coo=str(item.get("coo", "")),
            remark=str(item.get("remark", "")),
        )
        products.append(product)
    return products


def import_document_lines(path: str, config: dict) -> list:
    """从采购订单(PO)等文件中提取单据明细行，返回符合 make_doc_line 字段结构的 dict 列表（未保存）"""
    raw = _run_extraction(path, DOCUMENT_LINES_PROMPT, config)
    data = _parse_json_response(raw)
    if isinstance(data, dict):
        data = data.get("items", [])
    if not isinstance(data, list):
        raise AIImportError("AI 返回的明细信息格式不正确（应包含 items 数组）。")
    lines = []
    for item in data:
        if not isinstance(item, dict):
            continue
        quantity = _to_float(item.get("quantity", 0))
        unit_price = _to_float(item.get("unit_price", 0))
        amount = _to_float(item.get("amount", 0))
        remark = str(item.get("remark", ""))

        # 若源文件本身含有「金额」列，用 数量×单价 与其核对：
        # 注意乘法满足交换律，此核对无法判断数量与单价是否被互换填反，
        # 但能发现两者之一被识别成完全错误数值的情况，提醒用户核对。
        if amount and abs(quantity * unit_price - amount) > max(0.01, amount * 0.01):
            remark = (remark + " ⚠数量×单价与金额不符，请核对").strip()

        pseudo_product = {
            "id": "",
            "model_no": str(item.get("model_no", "")),
            "name_cn": str(item.get("name_cn", "")),
            "name_en": str(item.get("name_en", "")),
            "hs_code": str(item.get("hs_code", "")),
            "unit": str(item.get("unit", "pcs")) or "pcs",
            "unit_price": unit_price,
            "net_weight": 0.0,
            "gross_weight": 0.0,
            "length_mm": 0.0,
            "width_mm": 0.0,
            "height_mm": 0.0,
            "coo": str(item.get("coo", "")),
            "remark": remark,
        }
        line = make_doc_line(pseudo_product, quantity=quantity)
        lines.append(line)
    return lines


def _to_float(value) -> float:
    try:
        return float(value)
    except (ValueError, TypeError):
        return 0.0
